"""
Execution Report Generator - Comprehensive documentation for automation runs.

Generates detailed reports in multiple formats:
- Markdown (.md) - Full documentation with embedded images
- HTML (.html) - Interactive report with screenshots
- PDF (.pdf) - Professional printable documentation
- DOCX (.docx) - Editable Word document
- JSON (.json) - Machine-readable data

Each report includes:
- Step-by-step execution details
- Screenshots at each step
- Timing and performance metrics
- AI-generated summary and observations
- Error analysis for failed steps
"""

import base64
import json
import logging
import os
from dataclasses import dataclass, field, asdict
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional, TYPE_CHECKING

if TYPE_CHECKING:
    from llm_web_agent.interfaces.llm import ILLMProvider
    from llm_web_agent.interfaces.browser import IPage

logger = logging.getLogger(__name__)


@dataclass
class StepDetail:
    """Detailed information about a single execution step."""
    step_number: int
    action: str
    target: str
    status: str  # "success", "failed", "skipped"
    
    # Timing
    started_at: Optional[datetime] = None
    completed_at: Optional[datetime] = None
    duration_ms: float = 0
    
    # Resolution details
    locator_type: Optional[str] = None
    selector_used: Optional[str] = None
    alternatives_tried: List[str] = field(default_factory=list)
    
    # Screenshots
    screenshot_before: Optional[str] = None  # Path
    screenshot_after: Optional[str] = None   # Path
    
    # Value/Result
    value: Optional[str] = None
    extracted_data: Optional[str] = None
    
    # Error info
    error: Optional[str] = None
    error_screenshot: Optional[str] = None
    
    # AI observations (generated post-run)
    ai_observation: Optional[str] = None


@dataclass
class ExecutionReport:
    """Complete execution report data."""
    
    # Run identification
    run_id: str
    created_at: datetime
    
    # Goal and outcome
    goal: str
    success: bool
    
    # Timing
    started_at: datetime
    completed_at: Optional[datetime] = None
    duration_seconds: float = 0
    
    # Statistics
    steps_total: int = 0
    steps_completed: int = 0
    steps_failed: int = 0
    
    # Site info
    initial_url: Optional[str] = None
    final_url: Optional[str] = None
    framework_detected: Optional[str] = None
    
    # Steps
    steps: List[StepDetail] = field(default_factory=list)
    
    # AI-generated content
    ai_summary: Optional[str] = None
    ai_key_observations: List[str] = field(default_factory=list)
    ai_failure_analysis: Optional[str] = None
    ai_recommendations: List[str] = field(default_factory=list)
    
    # Metadata
    browser: str = "chromium"
    llm_provider: str = "unknown"
    agent_version: str = "1.0.0"
    
    def to_dict(self) -> Dict[str, Any]:
        """Convert to dictionary for serialization."""
        return {
            "run_id": self.run_id,
            "created_at": self.created_at.isoformat(),
            "goal": self.goal,
            "success": self.success,
            "started_at": self.started_at.isoformat(),
            "completed_at": self.completed_at.isoformat() if self.completed_at else None,
            "duration_seconds": self.duration_seconds,
            "steps_total": self.steps_total,
            "steps_completed": self.steps_completed,
            "steps_failed": self.steps_failed,
            "initial_url": self.initial_url,
            "final_url": self.final_url,
            "framework_detected": self.framework_detected,
            "steps": [
                {
                    "step_number": s.step_number,
                    "action": s.action,
                    "target": s.target,
                    "status": s.status,
                    "duration_ms": s.duration_ms,
                    "locator_type": s.locator_type,
                    "selector_used": s.selector_used,
                    "value": s.value,
                    "error": s.error,
                    "screenshot_before": s.screenshot_before,
                    "screenshot_after": s.screenshot_after,
                    "ai_observation": s.ai_observation,
                }
                for s in self.steps
            ],
            "ai_summary": self.ai_summary,
            "ai_key_observations": self.ai_key_observations,
            "ai_failure_analysis": self.ai_failure_analysis,
            "ai_recommendations": self.ai_recommendations,
            "browser": self.browser,
            "llm_provider": self.llm_provider,
            "agent_version": self.agent_version,
        }


class ExecutionReportGenerator:
    """
    Generate comprehensive execution reports in multiple formats.
    
    Usage:
        generator = ExecutionReportGenerator(output_dir="./reports")
        
        # After execution
        report = generator.create_report(engine_result, screenshots)
        
        # Generate AI summaries
        await generator.generate_ai_content(report, llm_provider)
        
        # Export to all formats
        generator.export_all(report)
    """
    
    def __init__(
        self,
        output_dir: str = "./reports",
        include_screenshots: bool = True,
        embed_images_in_html: bool = True,
    ):
        self.output_dir = Path(output_dir)
        self.include_screenshots = include_screenshots
        self.embed_images = embed_images_in_html
        
        # Ensure output directory exists
        self.output_dir.mkdir(parents=True, exist_ok=True)
    
    def create_report(
        self,
        run_id: str,
        goal: str,
        success: bool,
        step_results: List[Any],  # List of StepResult from AdaptiveEngine
        duration_seconds: float,
        framework_detected: Optional[str] = None,
        screenshots: Optional[Dict[int, str]] = None,
    ) -> ExecutionReport:
        """
        Create an ExecutionReport from AdaptiveEngine results.
        
        Args:
            run_id: Unique run identifier
            goal: Original goal
            success: Overall success
            step_results: List of StepResult objects
            duration_seconds: Total duration
            framework_detected: Detected framework
            screenshots: Map of step_number -> screenshot_path
        """
        now = datetime.now()
        screenshots = screenshots or {}
        
        steps = []
        for i, sr in enumerate(step_results, 1):
            step = StepDetail(
                step_number=i,
                action=sr.step.action.value if hasattr(sr.step, 'action') else 'unknown',
                target=sr.step.target if hasattr(sr.step, 'target') else '',
                status="success" if sr.success else "failed",
                duration_ms=sr.duration_ms,
                locator_type=sr.locator_type.value if sr.locator_type else None,
                selector_used=sr.selector_used,
                value=sr.step.value if hasattr(sr.step, 'value') else None,
                error=sr.error,
                screenshot_after=screenshots.get(i),
            )
            steps.append(step)
        
        return ExecutionReport(
            run_id=run_id,
            created_at=now,
            goal=goal,
            success=success,
            started_at=now,
            completed_at=now,
            duration_seconds=duration_seconds,
            steps_total=len(steps),
            steps_completed=sum(1 for s in steps if s.status == "success"),
            steps_failed=sum(1 for s in steps if s.status == "failed"),
            framework_detected=framework_detected,
            steps=steps,
        )
    
    async def generate_ai_content(
        self,
        report: ExecutionReport,
        llm_provider: "ILLMProvider",
    ) -> None:
        """
        Generate AI summary, observations, and failure analysis.
        
        Args:
            report: ExecutionReport to enhance
            llm_provider: LLM provider for generating content
        """
        from llm_web_agent.interfaces.llm import Message
        
        # Build context from report
        steps_summary = "\n".join([
            f"Step {s.step_number}: {s.action} on '{s.target}' - {s.status}"
            + (f" (error: {s.error})" if s.error else "")
            for s in report.steps
        ])
        
        prompt = f"""Analyze this browser automation execution and provide a structured summary.

## Goal
{report.goal}

## Outcome
- Success: {report.success}
- Duration: {report.duration_seconds:.1f}s
- Steps: {report.steps_completed}/{report.steps_total} completed

## Step Details
{steps_summary}

## Provide (respond in JSON format):
{{
  "summary": "2-3 sentence summary of what happened",
  "key_observations": ["observation 1", "observation 2", ...],
  "failure_analysis": "if any failures, explain why and how to fix (null if no failures)",
  "recommendations": ["suggestion for improvement 1", ...]
}}

Only output valid JSON, no markdown."""

        try:
            response = await llm_provider.complete(
                [Message.user(prompt)],
                temperature=0.3
            )
            
            # Parse response
            content = response.content.strip()
            if '```json' in content:
                content = content.split('```json')[1].split('```')[0]
            elif '```' in content:
                content = content.split('```')[1].split('```')[0]
            
            data = json.loads(content)
            
            report.ai_summary = data.get('summary')
            report.ai_key_observations = data.get('key_observations', [])
            report.ai_failure_analysis = data.get('failure_analysis')
            report.ai_recommendations = data.get('recommendations', [])
            
            logger.info("Generated AI content for report")
            
        except Exception as e:
            logger.error(f"Failed to generate AI content: {e}")
            report.ai_summary = f"AI summary generation failed: {e}"
    
    def export_json(self, report: ExecutionReport, filename: Optional[str] = None) -> Path:
        """Export report as JSON."""
        filename = filename or f"{report.run_id}_report.json"
        path = self.output_dir / filename
        
        with open(path, 'w') as f:
            json.dump(report.to_dict(), f, indent=2)
        
        logger.info(f"Exported JSON report: {path}")
        return path
    
    def export_markdown(self, report: ExecutionReport, filename: Optional[str] = None) -> Path:
        """Export report as Markdown."""
        filename = filename or f"{report.run_id}_report.md"
        path = self.output_dir / filename
        
        md = self._generate_markdown(report)
        path.write_text(md)
        
        logger.info(f"Exported Markdown report: {path}")
        return path
    
    def _generate_markdown(self, report: ExecutionReport) -> str:
        """Generate Markdown content."""
        status_emoji = "✅" if report.success else "❌"
        
        md = f"""# Execution Report: {report.run_id}

**Generated:** {report.created_at.strftime('%Y-%m-%d %H:%M:%S')}

---

## Summary

| Metric | Value |
|--------|-------|
| **Goal** | {report.goal[:80]}{'...' if len(report.goal) > 80 else ''} |
| **Status** | {status_emoji} {'Success' if report.success else 'Failed'} |
| **Duration** | {report.duration_seconds:.1f}s |
| **Steps** | {report.steps_completed}/{report.steps_total} completed |
| **Framework** | {report.framework_detected or 'N/A'} |

"""
        
        # AI Summary
        if report.ai_summary:
            md += f"""## AI Summary

{report.ai_summary}

"""
        
        # Key Observations
        if report.ai_key_observations:
            md += """## Key Observations

"""
            for obs in report.ai_key_observations:
                md += f"- {obs}\n"
            md += "\n"
        
        # Step Details
        md += """## Step-by-Step Details

"""
        for step in report.steps:
            status = "✅" if step.status == "success" else "❌"
            md += f"""### Step {step.step_number}: {step.action.upper()}

- **Target:** {step.target}
- **Status:** {status} {step.status}
- **Duration:** {step.duration_ms:.0f}ms
- **Locator:** {step.locator_type or 'N/A'}
"""
            if step.selector_used:
                md += f"- **Selector:** `{step.selector_used}`\n"
            if step.value:
                md += f"- **Value:** {step.value}\n"
            if step.error:
                md += f"- **Error:** {step.error}\n"
            if step.ai_observation:
                md += f"- **Observation:** {step.ai_observation}\n"
            
            # Screenshot
            if step.screenshot_after and self.include_screenshots:
                screenshot_path = Path(step.screenshot_after)
                if screenshot_path.exists():
                    md += f"\n![Step {step.step_number}]({screenshot_path.absolute()})\n"
            
            md += "\n"
        
        # Failure Analysis
        if report.ai_failure_analysis:
            md += f"""## Failure Analysis

{report.ai_failure_analysis}

"""
        
        # Recommendations
        if report.ai_recommendations:
            md += """## Recommendations

"""
            for rec in report.ai_recommendations:
                md += f"- {rec}\n"
            md += "\n"
        
        # Metadata
        md += f"""---

## Metadata

- **Run ID:** {report.run_id}
- **Browser:** {report.browser}
- **LLM Provider:** {report.llm_provider}
- **Agent Version:** {report.agent_version}
"""
        
        return md
    
    def export_html(self, report: ExecutionReport, filename: Optional[str] = None) -> Path:
        """Export report as HTML with embedded screenshots."""
        filename = filename or f"{report.run_id}_report.html"
        path = self.output_dir / filename
        
        html = self._generate_html(report)
        path.write_text(html)
        
        logger.info(f"Exported HTML report: {path}")
        return path
    
    def _generate_html(self, report: ExecutionReport) -> str:
        """Generate HTML content with embedded CSS."""
        status_class = "success" if report.success else "failed"
        status_text = "✅ Success" if report.success else "❌ Failed"
        
        steps_html = ""
        for step in report.steps:
            step_status = "success" if step.status == "success" else "failed"
            
            screenshot_html = ""
            if step.screenshot_after and self.include_screenshots:
                screenshot_path = Path(step.screenshot_after)
                if screenshot_path.exists():
                    if self.embed_images:
                        # Embed as base64
                        with open(screenshot_path, 'rb') as f:
                            img_data = base64.b64encode(f.read()).decode()
                        screenshot_html = f'<img src="data:image/png;base64,{img_data}" class="screenshot" alt="Step {step.step_number}">'
                    else:
                        screenshot_html = f'<img src="{screenshot_path.absolute()}" class="screenshot" alt="Step {step.step_number}">'
            
            steps_html += f"""
            <div class="step {step_status}">
                <h3>Step {step.step_number}: {step.action.upper()}</h3>
                <table>
                    <tr><td><strong>Target:</strong></td><td>{step.target}</td></tr>
                    <tr><td><strong>Status:</strong></td><td class="{step_status}">{step.status}</td></tr>
                    <tr><td><strong>Duration:</strong></td><td>{step.duration_ms:.0f}ms</td></tr>
                    <tr><td><strong>Locator:</strong></td><td>{step.locator_type or 'N/A'}</td></tr>
                    {f'<tr><td><strong>Selector:</strong></td><td><code>{step.selector_used}</code></td></tr>' if step.selector_used else ''}
                    {f'<tr><td><strong>Value:</strong></td><td>{step.value}</td></tr>' if step.value else ''}
                    {f'<tr><td><strong>Error:</strong></td><td class="error">{step.error}</td></tr>' if step.error else ''}
                </table>
                {screenshot_html}
            </div>
            """
        
        observations_html = ""
        if report.ai_key_observations:
            observations_html = "<ul>" + "".join(f"<li>{obs}</li>" for obs in report.ai_key_observations) + "</ul>"
        
        recommendations_html = ""
        if report.ai_recommendations:
            recommendations_html = "<ul>" + "".join(f"<li>{rec}</li>" for rec in report.ai_recommendations) + "</ul>"
        
        html = f"""<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Execution Report: {report.run_id}</title>
    <style>
        * {{ box-sizing: border-box; }}
        body {{
            font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif;
            line-height: 1.6;
            max-width: 1000px;
            margin: 0 auto;
            padding: 20px;
            background: #f5f5f5;
            color: #333;
        }}
        .header {{
            background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
            color: white;
            padding: 30px;
            border-radius: 10px;
            margin-bottom: 20px;
        }}
        .header h1 {{ margin: 0 0 10px 0; }}
        .summary-card {{
            background: white;
            padding: 20px;
            border-radius: 10px;
            margin-bottom: 20px;
            box-shadow: 0 2px 5px rgba(0,0,0,0.1);
        }}
        .summary-card h2 {{ margin-top: 0; color: #667eea; }}
        .status.success {{ color: #27ae60; font-weight: bold; }}
        .status.failed {{ color: #e74c3c; font-weight: bold; }}
        .step {{
            background: white;
            padding: 20px;
            border-radius: 10px;
            margin-bottom: 15px;
            box-shadow: 0 2px 5px rgba(0,0,0,0.1);
            border-left: 4px solid #ddd;
        }}
        .step.success {{ border-left-color: #27ae60; }}
        .step.failed {{ border-left-color: #e74c3c; }}
        .step h3 {{ margin-top: 0; color: #333; }}
        .step table {{ width: 100%; border-collapse: collapse; }}
        .step td {{ padding: 5px 10px 5px 0; }}
        .step td:first-child {{ width: 120px; color: #666; }}
        .screenshot {{
            max-width: 100%;
            border-radius: 5px;
            margin-top: 15px;
            box-shadow: 0 2px 10px rgba(0,0,0,0.2);
        }}
        code {{
            background: #f0f0f0;
            padding: 2px 6px;
            border-radius: 3px;
            font-size: 0.9em;
        }}
        .error {{ color: #e74c3c; }}
        .ai-section {{
            background: linear-gradient(135deg, #f8f9fa 0%, #e9ecef 100%);
            padding: 20px;
            border-radius: 10px;
            margin-bottom: 20px;
        }}
        .ai-section h2 {{ color: #764ba2; }}
        ul {{ margin: 10px 0; padding-left: 20px; }}
        .metadata {{
            color: #666;
            font-size: 0.9em;
            padding: 15px;
            background: #f8f9fa;
            border-radius: 5px;
        }}
    </style>
</head>
<body>
    <div class="header">
        <h1>🤖 Execution Report</h1>
        <p>Run ID: {report.run_id} | {report.created_at.strftime('%Y-%m-%d %H:%M:%S')}</p>
    </div>
    
    <div class="summary-card">
        <h2>Summary</h2>
        <table>
            <tr><td><strong>Goal:</strong></td><td>{report.goal}</td></tr>
            <tr><td><strong>Status:</strong></td><td class="status {status_class}">{status_text}</td></tr>
            <tr><td><strong>Duration:</strong></td><td>{report.duration_seconds:.1f}s</td></tr>
            <tr><td><strong>Steps:</strong></td><td>{report.steps_completed}/{report.steps_total} completed</td></tr>
            <tr><td><strong>Framework:</strong></td><td>{report.framework_detected or 'N/A'}</td></tr>
        </table>
    </div>
    
    {f'''<div class="ai-section">
        <h2>🧠 AI Summary</h2>
        <p>{report.ai_summary}</p>
    </div>''' if report.ai_summary else ''}
    
    {f'''<div class="ai-section">
        <h2>🔍 Key Observations</h2>
        {observations_html}
    </div>''' if report.ai_key_observations else ''}
    
    <div class="summary-card">
        <h2>Step-by-Step Details</h2>
    </div>
    
    {steps_html}
    
    {f'''<div class="ai-section">
        <h2>⚠️ Failure Analysis</h2>
        <p>{report.ai_failure_analysis}</p>
    </div>''' if report.ai_failure_analysis else ''}
    
    {f'''<div class="ai-section">
        <h2>💡 Recommendations</h2>
        {recommendations_html}
    </div>''' if report.ai_recommendations else ''}
    
    <div class="metadata">
        <strong>Metadata:</strong> 
        Browser: {report.browser} | 
        LLM: {report.llm_provider} | 
        Version: {report.agent_version}
    </div>
</body>
</html>"""
        
        return html
    
    def export_pdf(self, report: ExecutionReport, filename: Optional[str] = None) -> Path:
        """
        Export report as PDF.
        
        Requires weasyprint or pdfkit installed.
        Falls back to HTML-to-PDF conversion.
        """
        filename = filename or f"{report.run_id}_report.pdf"
        path = self.output_dir / filename
        
        # Generate HTML first
        html = self._generate_html(report)
        
        try:
            # Try weasyprint
            from weasyprint import HTML
            HTML(string=html).write_pdf(path)
            logger.info(f"Exported PDF report (weasyprint): {path}")
            return path
        except ImportError:
            pass
        
        try:
            # Try pdfkit
            import pdfkit
            pdfkit.from_string(html, str(path))
            logger.info(f"Exported PDF report (pdfkit): {path}")
            return path
        except ImportError:
            pass
        
        # Fallback: save HTML and note
        html_path = self.output_dir / f"{report.run_id}_report.html"
        html_path.write_text(html)
        logger.warning(f"PDF export requires weasyprint or pdfkit. Saved HTML instead: {html_path}")
        return html_path
    
    def export_docx(self, report: ExecutionReport, filename: Optional[str] = None) -> Path:
        """
        Export report as DOCX (Word document).
        
        Requires python-docx installed.
        """
        filename = filename or f"{report.run_id}_report.docx"
        path = self.output_dir / filename
        
        try:
            from docx import Document
            from docx.shared import Inches, Pt, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            
            doc = Document()
            
            # Title
            title = doc.add_heading(f'Execution Report: {report.run_id}', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            doc.add_paragraph(f"Generated: {report.created_at.strftime('%Y-%m-%d %H:%M:%S')}")
            doc.add_paragraph()
            
            # Summary
            doc.add_heading('Summary', level=1)
            summary_table = doc.add_table(rows=5, cols=2)
            summary_table.style = 'TableGrid'
            rows = summary_table.rows
            rows[0].cells[0].text = "Goal"
            rows[0].cells[1].text = report.goal
            rows[1].cells[0].text = "Status"
            rows[1].cells[1].text = "Success" if report.success else "Failed"
            rows[2].cells[0].text = "Duration"
            rows[2].cells[1].text = f"{report.duration_seconds:.1f}s"
            rows[3].cells[0].text = "Steps"
            rows[3].cells[1].text = f"{report.steps_completed}/{report.steps_total}"
            rows[4].cells[0].text = "Framework"
            rows[4].cells[1].text = report.framework_detected or "N/A"
            
            doc.add_paragraph()
            
            # AI Summary
            if report.ai_summary:
                doc.add_heading('AI Summary', level=1)
                doc.add_paragraph(report.ai_summary)
            
            # Key Observations
            if report.ai_key_observations:
                doc.add_heading('Key Observations', level=1)
                for obs in report.ai_key_observations:
                    doc.add_paragraph(obs, style='List Bullet')
            
            # Step Details
            doc.add_heading('Step-by-Step Details', level=1)
            for step in report.steps:
                doc.add_heading(f'Step {step.step_number}: {step.action.upper()}', level=2)
                
                p = doc.add_paragraph()
                p.add_run("Target: ").bold = True
                p.add_run(step.target)
                
                p = doc.add_paragraph()
                p.add_run("Status: ").bold = True
                status_run = p.add_run(step.status)
                if step.status == "failed":
                    status_run.font.color.rgb = RGBColor(231, 76, 60)
                
                p = doc.add_paragraph()
                p.add_run("Duration: ").bold = True
                p.add_run(f"{step.duration_ms:.0f}ms")
                
                if step.error:
                    p = doc.add_paragraph()
                    p.add_run("Error: ").bold = True
                    error_run = p.add_run(step.error)
                    error_run.font.color.rgb = RGBColor(231, 76, 60)
                
                # Add screenshot if exists
                if step.screenshot_after and self.include_screenshots:
                    screenshot_path = Path(step.screenshot_after)
                    if screenshot_path.exists():
                        doc.add_picture(str(screenshot_path), width=Inches(6))
                
                doc.add_paragraph()
            
            # Failure Analysis
            if report.ai_failure_analysis:
                doc.add_heading('Failure Analysis', level=1)
                doc.add_paragraph(report.ai_failure_analysis)
            
            # Recommendations
            if report.ai_recommendations:
                doc.add_heading('Recommendations', level=1)
                for rec in report.ai_recommendations:
                    doc.add_paragraph(rec, style='List Bullet')
            
            doc.save(path)
            logger.info(f"Exported DOCX report: {path}")
            return path
            
        except ImportError:
            logger.warning("DOCX export requires python-docx. Install with: pip install python-docx")
            # Fallback to markdown
            md_path = self.export_markdown(report, f"{report.run_id}_report.md")
            return md_path
    
    def export_all(self, report: ExecutionReport, formats: List[str] = None) -> Dict[str, Path]:
        """
        Export report to all requested formats.
        
        Args:
            report: ExecutionReport to export
            formats: List of formats ('json', 'md', 'html', 'pdf', 'docx')
                     Default: all formats
        
        Returns:
            Dict mapping format to file path
        """
        formats = formats or ['json', 'md', 'html', 'pdf', 'docx']
        results = {}
        
        for fmt in formats:
            try:
                if fmt == 'json':
                    results['json'] = self.export_json(report)
                elif fmt == 'md':
                    results['md'] = self.export_markdown(report)
                elif fmt == 'html':
                    results['html'] = self.export_html(report)
                elif fmt == 'pdf':
                    results['pdf'] = self.export_pdf(report)
                elif fmt == 'docx':
                    results['docx'] = self.export_docx(report)
            except Exception as e:
                logger.error(f"Failed to export {fmt}: {e}")
        
        return results
